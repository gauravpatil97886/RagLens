"""Get plain text out of an uploaded file. No OCR, no layout analysis — just text."""

from __future__ import annotations

import io
from pathlib import Path

import docx
from pypdf import PdfReader


class ExtractionError(Exception):
    """Raised when a file is unsupported or contains no extractable text."""


ALLOWED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}

MIME_TYPES = {
    ".pdf": "application/pdf",
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def extension_of(filename: str) -> str:
    return Path(filename or "").suffix.lower()


def mime_for(filename: str) -> str:
    return MIME_TYPES.get(extension_of(filename), "application/octet-stream")


def extract_text(filename: str, data: bytes) -> str:
    ext = extension_of(filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise ExtractionError(
            f"Unsupported file type '{ext or filename}'. "
            f"Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )

    if ext == ".pdf":
        text = _from_pdf(data, filename)
    elif ext == ".docx":
        text = _from_docx(data, filename)
    else:
        text = data.decode("utf-8", errors="replace")

    text = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    if not text:
        raise ExtractionError(
            f"No text could be extracted from '{filename}'. If this is a scanned or "
            "image-only PDF it has no text layer, and this demo does not run OCR — "
            "re-export it as a text PDF or upload a .txt/.md/.docx version."
        )
    return text


def _from_pdf(data: bytes, filename: str) -> str:
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(f"'{filename}' could not be opened as a PDF: {exc}") from exc
    if reader.is_encrypted:
        # An empty user password is common; try it before giving up.
        try:
            reader.decrypt("")
        except Exception as exc:
            raise ExtractionError(f"'{filename}' is password-protected.") from exc
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")  # one bad page should not fail the whole document
    return "\n\n".join(p.strip() for p in pages if p.strip())


def _from_docx(data: bytes, filename: str) -> str:
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(f"'{filename}' could not be opened as a .docx: {exc}") from exc
    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)
